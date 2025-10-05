from docx import Document
from docx.shared import Inches
import os
from langchain_ollama import OllamaLLM

import base64
from dotenv import load_dotenv, find_dotenv
load_dotenv(find_dotenv())
URL_LLM = os.getenv("URL_LLM")
USER_LLM = os.getenv("USER_LLM")
PASSWORD_LLM = os.getenv("PASSWORD_LLM")


from docx.shared import Inches
import os
import json

def extract_elements_with_images(file_path):
    """
    Извлекает элементы документа в порядке появления: текст и изображения.
    Возвращает список: [{'type': 'text', 'text': ...}, {'type': 'image', 'blob': ..., 'ext': ...}]
    """
    doc = Document(file_path)
    elements = []

    # Собираем все связи (включая изображения)
    rels = doc.part.rels

    # Проходим по каждому элементу в теле документа
    for element in doc.element.body:
        if element.tag.endswith('p'):  # абзац
            para = Document(doc.element).paragraphs[-1]  # временный способ получить текст
            # Но лучше: извлечь текст напрямую из XML
            p_text = ""
            for node in element.iter():
                if node.tag.endswith('t'):  # w:t — текст
                    if node.text:
                        p_text += node.text
            p_text = p_text.strip()
            if p_text:
                elements.append({'type': 'text', 'text': p_text})

        elif element.tag.endswith('tbl'):  # таблица
            # Упрощённо: извлечём весь текст из таблицы
            table_text = ""
            for row in element.iter():
                if row.tag.endswith('t'):
                    if row.text:
                        table_text += row.text + " "
            table_text = table_text.strip()
            if table_text:
                elements.append({'type': 'text', 'text': table_text})

        # Проверка на изображение (в drawing или inline)
        elif any("drawing" in child.tag for child in element.iter()):
            for drawing in element.iter():
                if "drawing" in drawing.tag:
                    for img_part in drawing.iter():
                        if "blip" in img_part.tag:
                            # Извлекаем rId
                            embed_attr = [at for at in img_part.attrib.values() if "embed" in at]
                            if embed_attr:
                                rId = embed_attr[0].split('}')[1] if '}' in embed_attr[0] else embed_attr[0]
                                if rId in rels:
                                    image_part = rels[rId].target_part
                                    image_blob = image_part.blob
                                    ext = image_part.content_type.split('/')[-1]
                                    img_filename = f"temp_img_{len([e for e in elements if e['type']=='image'])}.{ext}"
                                    with open(img_filename, "wb") as f:
                                        f.write(image_blob)
                                    elements.append({
                                        'type': 'image',
                                        'filename': img_filename
                                    })

    return elements


def create_prompt_for_headings(text_blocks):
    """Генерация промта для LLM"""
    text_preview = "\n\n".join([f"Блок {i+1}: {block[:300]}..." for i, block in enumerate(text_blocks[:10])])
    prompt = f"""
    Проанализируй следующие текстовые блоки и разбей их на логические разделы.
    Для каждого раздела укажи:
    - start_block: индекс первого блока (с 0)
    - end_block: индекс последнего блока
    - title: название раздела

    Верни ТОЛЬКО массив в формате JSON:

    [
    {{
        "start_block": 0,
        "end_block": 4,
        "title": "Введение"
    }},
    {{
        "start_block": 5,
        "end_block": 8,
        "title": "Методология"
    }}
    ]

    Текстовые блоки:
    {text_preview}
    """
    return prompt


def parse_json_response(response):
    """Парсинг JSON из ответа модели"""
    try:
        start = response.find('[')
        end = response.rfind(']') + 1
        if start == -1 or end == 0:
            return []
        json_str = response[start:end]
        return json.loads(json_str)
    except Exception as e:
        print("Ошибка при парсинге JSON:", e)
        return []


def create_document_with_headings(elements, section_info, output_path):
    """Создаёт новый документ с заголовками и сохраняет изображения"""
    doc = Document()

    text_block_index = 0

    for elem in elements:
        if elem['type'] == 'text':
            # Проверяем, нужно ли вставить заголовок перед этим блоком
            if section_info and text_block_index == section_info[0]['start_block']:
                doc.add_heading(section_info[0]['title'], level=1)
                # Убираем обработанный раздел
                section_info = section_info[1:]

            doc.add_paragraph(elem['text'])
            text_block_index += 1

        elif elem['type'] == 'image':
            try:
                doc.add_picture(elem['filename'], width=Inches(5))
                # Можно добавить пустой абзац после картинки для отступа
                doc.add_paragraph()
            except Exception as e:
                print(f"Ошибка вставки изображения {elem['filename']}: {e}")

    doc.save(output_path)
    print(f"✅ Новый документ сохранён: {output_path}")


def cleanup_temp_images(elements):
    """Удаляет временные изображения"""
    for elem in elements:
        if elem['type'] == 'image' and os.path.exists(elem['filename']):
            os.remove(elem['filename'])


# === ОСНОВНОЙ КОД ===
if __name__ == "__main__":
    input_file = "/Users/alexeyvaganov/Documents/Project/pro_club_2/pro_club_season_2/prep/razdel/test.docx"
    output_file = "output_with_headings.docx"

    print("Чтение элементов документа...")
    elements = extract_elements_with_images(input_file)

    # Соберём только текстовые блоки для анализа
    text_blocks = [el['text'] for el in elements if el['type'] == 'text']

    if not text_blocks:
        print("Нет текста для анализа.")
    else:
        print(f"Найдено {len(text_blocks)} текстовых блоков.")

        # Генерируем промт
        prompt = create_prompt_for_headings(text_blocks)

        # Отправляем в LLM
        print("Отправка в модель для определения разделов...")
        encoded_credentials = base64.b64encode(f"{USER_LLM}:{PASSWORD_LLM}".encode()).decode()
        headers = {'Authorization': f'Basic {encoded_credentials}'}

        llm_class = OllamaLLM(model="gemma3:12b", temperature = 0.1, base_url=URL_LLM, client_kwargs={'headers': headers})
        response = llm_class.invoke(prompt)

        print("Ответ модели:\n", response)

        # Парсим ответ
        section_info = parse_json_response(response)

        if not section_info:
            print("Не удалось определить разделы. Создаю без заголовков.")
            section_info = []

        # Создаём новый документ
        create_document_with_headings(elements, section_info, output_file)

    # Очистка временных файлов
    cleanup_temp_images(elements)