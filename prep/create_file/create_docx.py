# pip install python-docx opencv-python

from docx import Document
from docx.shared import Mm
import os
import sys
import json
import re

from dotenv import load_dotenv, find_dotenv
load_dotenv(find_dotenv())
URL_LLM = os.getenv("URL_LLM")
USER_LLM = os.getenv("USER_LLM")
PASSWORD_LLM = os.getenv("PASSWORD_LLM") 
MODEL = os.getenv("MODEL")

# Добавляем родительский каталог в пути поиска модулей
current_dir = os.path.dirname(os.path.abspath(__file__))
parent_dir = os.path.dirname(current_dir)
sys.path.insert(0, parent_dir)

# Импорты из соседних модулей
from text_to_paragraphs.text_to_paragraphs import text_to_paragraphs
from picture_description.picture_description import picture_description
from text_modifier.text_modifier import TextModify

# === LLM для разбиения на разделы ===
from langchain_ollama import OllamaLLM
import base64

encoded_credentials = base64.b64encode(f"{USER_LLM}:{PASSWORD_LLM}".encode()).decode()
headers = {'Authorization': f'Basic {encoded_credentials}'}


def image_is_required(paragraph):

    llm = OllamaLLM(model=MODEL, temperature=0.1, base_url=URL_LLM, client_kwargs={'headers': headers})
    # prompt = "Тебе необходимо определить требует ли текст добавления картинки. " \
    # "Необходимо ориентироваться на слова (или их аналоги, догадывайся по смыслу): показать, отбор, перейти, посмотрите, сейчас на экране, открыть, выберем, нажмем, нажать .  " \
    # "Если нужна картинка ответь 1 если не нужна ответь 0. Только результат без пояснений. " \
    # "Вот текст: " + paragraph
    prompt = (
    "Ты — эксперт по технической документации. Определи, требует ли данный фрагмент инструкции вставки иллюстрации (скриншота, схемы, фото интерфейса и т.п.).\n"
    "Придерживайся следующих правил, иллюстрация нужна, если в тексте:\n"
    "- описывается конкретный элемент интерфейса (кнопка, меню, поле ввода и т.д.),\n"
    "- даётся пошаговое руководство с действиями пользователя (нажать, выбрать, перейти, открыть, следует перейти, необходимо перейти  и т.п.),\n"
    "- упоминается визуальный результат (вы увидите, на экране появится, как показано на рисунке, в данном элементе),\n"
    "- есть ссылка на расположение чего-либо (в левом верхнем углу, под заголовком, она заполняется  и т.п.).\n\n"
    "- при необходимости догадывайся по контексту.\n"
    "Если иллюстрация **полезна или необходима** для понимания — ответь «1».\n"
    "Если не нужна — ответь «0».\n\n"
    "Ответ должен содержать **только одну цифру**: 1 или 0. Без пояснений, без знаков препинания.\n\n"
    "Текст: " + paragraph
    )
    try:
        response = llm.invoke(prompt)
    except Exception as e:
        print(f"Ошибка при вызове LLM: {e}")
        
    return response

def table_segments_time(json_file_path):
    with open(json_file_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    rows = []
    buffer = ""  # сюда складываем "висящее" начало предложения

    for i, seg in enumerate(data.get("segments", [])):
        text = seg["text"].strip()

        # добавляем "хвост" из предыдущего сегмента, если он был
        if buffer:
            text = buffer + " " + text
            buffer = ""

        # режем на предложения по ". "
        sentences = text.split(". ")

        # проверяем, закончился ли последний кусок точкой
        last_part = sentences[-1]
        if not last_part.endswith((".", "!", "?")):
            # предложение не закончилось → переносим в buffer для следующего сегмента
            buffer = sentences.pop()

        # теперь фиксируем завершённые предложения
        for sent in sentences:
            sent = sent.strip()
            if sent:
                if not sent.endswith((".", "!", "?")):
                    sent += "."
                rows.append((sent, seg["end"]))

    # на всякий случай — если файл закончился, а buffer остался
    if buffer:
        rows.append((buffer, data["segments"][-1]["end"]))

    return rows

def get_sections_from_llm(paragraphs, max_paragraphs_per_chunk=20):
    """
    Разбивает список абзацев на куски и отправляет каждый в LLM для определения начала разделов.
    Возвращает список разделов в виде:
    [{'title': str, 'start_par': int}]
    Границы разделов определяются как:
      - от start_par текущего раздела
      - до start_par следующего раздела (исключительно)
      - или до конца текста, если это последний раздел.
    Это гарантирует, что все абзацы будут включены без пропусков.
    """
    llm = OllamaLLM(model=MODEL, temperature=0.1, base_url=URL_LLM, client_kwargs={'headers': headers})

    # Собираем все найденные точки начала разделов: {номер_абзаца: название}
    section_starts = {}

    # === Обрабатываем текст по частям ===
    for chunk_start in range(0, len(paragraphs), max_paragraphs_per_chunk):
        chunk = paragraphs[chunk_start:chunk_start + max_paragraphs_per_chunk]
        numbered_chunk = "\n".join([f"[{chunk_start + i + 1}] {p}" for i, p in enumerate(chunk)])

        # prompt = f"""
        # Проанализируй следующие пронумерованные абзацы и определи, с каких абзацев начинаются новые логические разделы.
        # ВАЖНО:
        # - Первый абзац всегда считается началом первого раздела.
        # - Укажи ТОЛЬКО абзацы, с которых начинается новый раздел.
        # - Не указывай конечные номера — только начала.

        # Формат ответа — строго по одному на строку:
        # [Номер] [Название раздела]

        # Пример:
        # 1 Введение
        # 5 Основной анализ
        # 12 Заключение

        # Текст:
        # {numbered_chunk}
        # """
        prompt = f"""
        Ты — эксперт по технической документации. Проанализируй пронумерованные абзацы инструкции и определи, с каких абзацев начинаются **новые логические разделы**.

        Разделы в инструкциях обычно включают:
        - Введение / Обзор
        - Требования / Предварительные условия
        - Пошаговые действия (например: «Шаг 1», «Настройка», «Запуск»)
        - Проверка результата / Верификация
        - Устранение неполадок
        - Заключение / Дополнительные рекомендации

        ВАЖНО:
        - Первый абзац ВСЕГДА является началом первого раздела.
        - Раздел начинается, если:
        • Тема или задача меняется;
        • Начинается новый этап инструкции;
        • Появляется подзаголовок по смыслу (даже если он не выделен явно).
        - Не выдумывай разделы — опирайся только на содержание текста.
        - Не включай промежуточные пояснения, примеры или примечания как отдельные разделы, если они не начинают новую тему.

        Формат ответа — строго по одной строке на раздел:
        <номер_абзаца> <название_раздела>

        Требования к формату:
        - Номер — целое число без скобок.
        - Название — краткое (3–6 слов), отражающее суть раздела, в стиле технической инструкции.
        - Никаких дополнительных символов, пояснений или пустых строк.

        Пример корректного ответа:
        1 Введение
        4 Предварительные требования
        7 Шаг 1: Запуск приложения
        12 Шаг 2: Настройка параметров
        18 Проверка результата
        22 Заключение

        Текст:
        {numbered_chunk}
        """

        try:
            response = llm.invoke(prompt)
        except Exception as e:
            print(f"Ошибка при вызове LLM для чанка [{chunk_start+1}–{chunk_start+len(chunk)}]: {e}")
            # Даже при ошибке считаем, что начало чанка — новая секция (на всякий случай)
            first_in_chunk = chunk_start + 1
            if first_in_chunk not in section_starts:
                section_starts[first_in_chunk] = "Раздел"
            continue

        # Парсим ответ: ищем строки вида "3 Название" или "[3] Название"
        lines = response.strip().splitlines()
        for line in lines:
            line = line.strip()
            if not line:
                continue
            # Поддерживаем форматы: "3 Название", "[3] Название", "3. Название"
            match = re.match(r"[\[\(]?\s*(\d+)\s*[\]\)]?\s*(.+)", line)
            if match:
                try:
                    par_num = int(match.group(1))
                    title = match.group(2).strip(" \"'.")
                    # Проверяем, что номер в пределах текущего чанка
                    if chunk_start + 1 <= par_num <= chunk_start + len(chunk):
                        section_starts[par_num] = title
                except ValueError:
                    continue

    # === Гарантируем, что первый абзац всегда начало раздела ===
    # if 1 not in section_starts:
    #     section_starts[1] = "Введение"

    # === Преобразуем в отсортированный список разделов ===
    sorted_starts = sorted(section_starts.items())  # [(1, "Введение"), (5, "Анализ"), ...]

    sections = []
    for i, (start_par, title) in enumerate(sorted_starts):
        sections.append({
            'title': title,
            'start_par': start_par
        })

    return sections


class create_docx:
    def __init__(self, json_file_path, video_path="", UseTextModify=False):
        self.json_file_path = json_file_path
        self.video_path = video_path
        self.UseTextModify = UseTextModify

    def get_docx(self):
        json_file_path = self.json_file_path
        UseTextModify = self.UseTextModify

        if not os.path.isfile(json_file_path):
            raise FileNotFoundError(f"Файл {json_file_path} не найден.")

        doc = Document()
        with open(json_file_path, 'r', encoding='utf-8') as file:
            data = json.load(file)

        full_text = data.get('full_text', '')

        # === Шаг 1: Поиск "сейчас на экране" в segments ===
        # table_time_screen = []
        # counter = 1
        # search_sequence = ["сейчас", "на", "экране"]
        # for segment in data['segments']:
        #     words = segment.get('words', [])
        #     lower_words = [word['word'].lower() for word in words]
        #     for i in range(len(lower_words) - 2):
        #         if lower_words[i:i+3] == search_sequence:
        #             start_time = words[i]['start']
        #             table_time_screen.append({
        #                 "Number": counter,
        #                 "Time": start_time
        #             })
        #             counter += 1
        #             break

        # === Шаг 2: Разбиваем текст на абзацы ===
        
        #1
        segments_time = table_segments_time(json_file_path)
        class_text_to_paragraphs = text_to_paragraphs(full_text, segments_time)
        paragraphs_table = class_text_to_paragraphs.get_text_to_paragraphs_table()        
        paragraphs = [p[0] for p in paragraphs_table]

        paragraphs_time_scr = {}

        for idx, row in enumerate(paragraphs_table, start=1):
            paragraph, end_time = row  # ("текст", end)

            image_is_required_result = image_is_required(paragraph)

            #if image_is_required_result == "1\n" or image_is_required_result == "1" or image_is_required_result == 1:
            if '1' in str(image_is_required_result):
                # Проверяем, есть ли уже такой end_time в словаре
                existing_keys = [k for k, v in paragraphs_time_scr.items() if v == end_time]

                if existing_keys:
                    # Если уже есть, удаляем старый ключ и вставляем новый
                    old_key = existing_keys[0]
                    del paragraphs_time_scr[old_key]
                    paragraphs_time_scr[idx] = end_time
                else:
                    # Если нет — добавляем
                    paragraphs_time_scr[idx] = end_time

        if UseTextModify==True:
            print("Проводим улучшение текста...")
            text_modifier = TextModify()
            for i in range(len(paragraphs)):
                paragraphs[i] = text_modifier.improve_text(paragraphs[i])
            

        # === Шаг 3: LLM разбивает на разделы (сохраняем оригинальные абзацы) ===
        print("Отправляем текст в LLM для разбиения на разделы...")
        sections = get_sections_from_llm(paragraphs)

        # Преобразуем в полные диапазоны без пропусков
        if sections:
            sections.sort(key=lambda x: x['start_par'])
            starts = [s['start_par'] for s in sections]
            starts.append(len(paragraphs) + 1)
            full_sections = []
            for i, sec in enumerate(sections):
                full_sections.append({
                    'title': sec['title'],
                    'start_par': sec['start_par'],
                    'end_par': min(starts[i + 1] - 1, len(paragraphs))
                })
            sections = full_sections
        else:
            # fallback: один раздел на весь текст
            sections = [{'title': 'Документ', 'start_par': 1, 'end_par': len(paragraphs)}]

        # === Шаг 4: Формируем документ с разделами и картинками ===
        video_path = self.video_path
        current_paragraph_index = 0  # Считаем, сколько абзацев текста уже вставлено

        if video_path != "" and os.path.isfile(video_path):
            class_picture_description = picture_description()

            # --- Режим A: Есть упоминания "сейчас на экране" ---
            if paragraphs_time_scr:
                print("Вставляем текст с разделами и кадрами по упоминаниям.")
                for section in sections:
                    # Вставляем заголовок раздела
                    doc.add_heading(section['title'], level=1)

                    # Вставляем абзацы этого раздела
                    for par_num in range(section['start_par'], section['end_par'] + 1):
                        if par_num <= len(paragraphs):
                            current_paragraph_index += 1
                            para_text = paragraphs[par_num - 1]
                            doc.add_paragraph('\t' + para_text)

                            # Проверяем, нужно ли вставить картинку ПОСЛЕ этого абзаца
                            if current_paragraph_index in paragraphs_time_scr:
                                time_screen = paragraphs_time_scr[current_paragraph_index]
                                total_seconds = int(time_screen)
                                hours = total_seconds // 3600
                                minutes = (total_seconds % 3600) // 60
                                seconds = total_seconds % 60
                                formatted_time = f"{hours:02}:{minutes:02}:{seconds:02}"
                                frame_at_time = class_picture_description.save_frame_at_time(
                                    video_path, formatted_time
                                )
                                doc.add_picture(frame_at_time, width=Mm(165))

            # --- Режим B: Нет упоминаний — равномерные кадры ---
            else:
                print("Нет упоминаний. Добавляем 5 равномерных кадров.")
                try:
                    total_duration = data['segments'][-1]['end']
                except (IndexError, KeyError, TypeError):
                    raise ValueError("Не удалось определить длительность видео.")

                time_stamps = [total_duration * i / 6 for i in range(1, 6)]
                frame_paths = []
                import uuid
                import shutil

                base_temp_name = "frame.jpg"
                video_dir = os.path.dirname(video_path)
                for i, t in enumerate(time_stamps):
                    total_seconds = int(t)
                    hours = total_seconds // 3600
                    minutes = (total_seconds % 3600) // 60
                    seconds = total_seconds % 60
                    formatted_time = f"{hours:02}:{minutes:02}:{seconds:02}"
                    temp_frame = class_picture_description.save_frame_at_time(video_path, formatted_time)
                    unique_path = f"temp_frame_{i}_{uuid.uuid4().hex[:8]}.jpg"
                    output_path = os.path.join(video_dir, unique_path)
                    shutil.copy(temp_frame, output_path)
                    frame_paths.append(output_path)

                # Определяем, после каких **общих** абзацев вставлять картинки
                total_paragraphs = len(paragraphs)
                image_positions = []
                num_images = len(frame_paths)
                for i in range(1, num_images + 1):
                    pos = int(round((i / (num_images + 1)) * total_paragraphs))
                    pos = max(1, min(pos, total_paragraphs))
                    image_positions.append(pos)

                # Вставляем разделы и картинки
                for section in sections:
                    doc.add_heading(section['title'], level=1)
                    for par_num in range(section['start_par'], section['end_par'] + 1):
                        if par_num <= len(paragraphs):
                            current_paragraph_index += 1
                            para_text = paragraphs[par_num - 1]
                            doc.add_paragraph('\t' + para_text)

                            if current_paragraph_index in image_positions:
                                img_idx = image_positions.index(current_paragraph_index)
                                doc.add_picture(frame_paths[img_idx], width=Mm(165))

        else:
            # Режим C: Только текст
            for section in sections:
                doc.add_heading(section['title'], level=1)
                for par_num in range(section['start_par'], section['end_par'] + 1):
                    if par_num <= len(paragraphs):
                        para_text = paragraphs[par_num - 1]
                        doc.add_paragraph('\t' + para_text)

        # === Шаг 5: Сохранение ===
        docx_file_path = os.path.splitext(json_file_path)[0] + '.docx'
        doc.save(docx_file_path)
        print(f"Документ с разделами сохранён: {docx_file_path}")
        return docx_file_path